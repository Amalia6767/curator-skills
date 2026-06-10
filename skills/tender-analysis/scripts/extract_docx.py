#!/usr/bin/env python3
"""
提取 Word 文档内容
用于投标文件分析
"""
import sys
from docx import Document

def extract_docx(file_path):
    """提取 Word 文档的所有文本内容"""
    try:
        doc = Document(file_path)
        content = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    content = extract_docx(file_path)
    
    if content:
        print(content)
    else:
        sys.exit(1)
