#!/usr/bin/env python3
"""
提取 .doc 和 .docx 文档内容
支持老版本的 .doc 格式和新版本的 .docx 格式
"""
import sys
import os

def extract_doc(file_path):
    """提取 .doc 或 .docx 文档内容"""
    try:
        # 尝试使用 docx2txt（支持 .doc 和 .docx）
        try:
            import docx2txt
            content = docx2txt.process(file_path)
            if content and content.strip():
                return content
        except ImportError:
            pass
        except Exception as e:
            print(f"docx2txt 处理失败: {e}", file=sys.stderr)
        
        # 如果是 .docx，尝试使用 python-docx
        if file_path.lower().endswith('.docx'):
            try:
                from docx import Document
                doc = Document(file_path)
                content = []
                
                # 提取段落文本
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content.append(" | ".join(row_text))
                
                return "\n".join(content)
            except Exception as e:
                print(f"python-docx 处理失败: {e}", file=sys.stderr)
        
        # 如果是 .doc，尝试使用系统工具
        if file_path.lower().endswith('.doc'):
            # 尝试使用 textutil (macOS)
            if sys.platform == 'darwin':
                import subprocess
                try:
                    result = subprocess.run(
                        ['textutil', '-convert', 'txt', '-stdout', file_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout:
                        return result.stdout
                except Exception as e:
                    print(f"textutil 处理失败: {e}", file=sys.stderr)
            
            return None
            
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract_doc.py <文件路径>")
        print("支持格式: .doc, .docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    content = extract_doc(file_path)
    
    if content:
        print(content)
    else:
        print(f"警告: 无法提取文件内容: {file_path}", file=sys.stderr)
        print("建议: 尝试将 .doc 文件转换为 .docx 格式", file=sys.stderr)
        sys.exit(1)
